from docx import Document

document = Document()
document.add_heading("Xin Nghi Viec", 0)

p = document.add_paragraph('Ho Ten: ')
p.add_run("NGuyen Van A").boid = True
p.add_run("Chuc Vu: Nhan Vien").italic = True

p = document.add_paragraph('Thoi Gian: ')
p.add_run("8AM ngay 15/11/2022").boid = True

document.add_heading('Noi Dung', level=1)
document.add_paragraph("Cong viec khong dung nang luc")

document.add_paragraph("Ki ten:")
document.add_paragraph("Nguyen Van A").italic = True

document.save("nghi_viec.docx")